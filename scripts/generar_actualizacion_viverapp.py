# -*- coding: utf-8 -*-
"""
Genera `actualizacion_viverapp.docx` a partir de `Guia_Usuario_ViverApp.docx`,
preservando todo el contenido original y añadiendo:
  - Un banner de "Novedades" al inicio.
  - Una sección final "8. Novedades de esta actualización" con detalle visual
    de todos los cambios añadidos a ViverApp en las últimas iteraciones.
"""
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from pathlib import Path

# =============================================================================
# RUTAS
# =============================================================================
SRC = Path(r"C:\Users\ivanf\Desktop\2026\EMPRESAS\AYTO SANTA CRUZ\Guia_Usuario_ViverApp.docx")
DST = Path(r"C:\Users\ivanf\Desktop\2026\EMPRESAS\AYTO SANTA CRUZ\actualizacion_viverapp.docx")

# Paleta de tonos (alineada con la UI de la app)
COLOR_AZUL_BG = "DBE7FD"      # Origen/destino
COLOR_AZUL_FG = "1D4ED8"
COLOR_VERDE_BG = "D6F5E6"     # Producto
COLOR_VERDE_FG = "065F46"
COLOR_AMBAR_BG = "FDE9C2"     # Detalles
COLOR_AMBAR_FG = "92400E"
COLOR_GRIS_BG = "F1F5F9"
COLOR_NEGRO = "0F172A"

# =============================================================================
# HELPERS
# =============================================================================

def set_cell_shading(cell, hex_color):
    """Sombrea una celda con el color de fondo dado (hex sin #)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_borders(cell, color="64748B", sz="8"):
    """Añade un borde fino a la celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def styled_paragraph(doc, text, bold=False, size=11, color=None, align=None, after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(after)
    return p


def callout_box(doc, title, body, bg_hex, fg_hex):
    """Renderiza una caja destacada estilo aviso (cabecera + cuerpo)."""
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.4)

    head = table.rows[0].cells[0]
    set_cell_shading(head, bg_hex)
    add_borders(head, color=fg_hex, sz="12")
    p = head.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(fg_hex)

    body_cell = table.rows[1].cells[0]
    set_cell_shading(body_cell, "FFFFFF")
    add_borders(body_cell, color=fg_hex, sz="12")
    bp = body_cell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(4)
    bp.paragraph_format.space_after = Pt(4)
    br = bp.add_run(body)
    br.font.size = Pt(10.5)
    br.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)

    doc.add_paragraph()  # separador


def add_bullet(doc, text, size=10.5):
    """Viñeta manual (el docx original solo tiene 'List Paragraph', sin
    autobullet). Anteponemos un "•" y aplicamos sangría."""
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"•  {text}")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
    return p


def add_section_divider(doc):
    """Línea separadora delgada en gris."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_kv_table(doc, header, rows, header_bg=COLOR_AZUL_BG, header_fg=COLOR_AZUL_FG):
    """Tabla simple con cabecera coloreada y filas alternas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    # (sin style: aplicamos bordes y sombreado manualmente)
    table.autofit = True

    # Cabecera
    for i, h in enumerate(header):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        add_borders(c, color=header_fg, sz="6")
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(header_fg)

    # Filas
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i + 1].cells[j]
            if i % 2 == 1:
                set_cell_shading(c, "F8FAFC")
            add_borders(c, color="CBD5E1", sz="4")
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10.5)
    doc.add_paragraph()


# =============================================================================
# CARGA DEL DOC ORIGINAL
# =============================================================================
print(f"Cargando {SRC}…")
doc = Document(str(SRC))

# =============================================================================
# 1) BANNER DE NOVEDADES AL INICIO
# =============================================================================
# Lo insertamos justo después del título de la guía (antes del índice).
# Para mantenerlo simple, lo añadimos antes del primer "Heading 1".
print("Insertando banner de novedades al inicio…")

# Localizamos el primer Heading 1 ("1. Introducción") y lo desplazamos
body = doc.element.body
first_h1 = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1":
        first_h1 = p._p
        break

# Construimos el banner como un sub-doc en memoria que luego prepende
def build_banner_paragraphs(doc):
    """Crea los párrafos del banner añadiéndolos al final del doc y devuelve
    los elementos XML para luego moverlos al inicio."""
    created = []

    # Título del banner
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("✨ Novedades de esta actualización")
    r1.bold = True
    r1.font.size = Pt(18)
    r1.font.color.rgb = RGBColor.from_string(COLOR_VERDE_FG)
    created.append(p1._p)

    # Subtítulo
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "Esta versión incorpora cambios importantes en Movimientos, Productos, "
        "Mapa del vivero y la primera experiencia de bienvenida. "
        "El detalle completo está en la sección 8 al final del documento."
    )
    r2.italic = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string(COLOR_NEGRO)
    created.append(p2._p)

    # Tabla resumen
    headers = ["Módulo", "Novedad principal", "Sección"]
    rows = [
        ("Bienvenida", "Modal de bienvenida con guía PDF y vídeo (solo primera vez).", "8.1"),
        ("Categorías de producto", "Cada categoría usa formato y unidad propia (kg, lt, m³, m, ud).", "8.2"),
        ("Mapa del vivero", "Nuevas zonas: Almacén Fitosanitarios, Almacén General, Almacén Fertilizantes y Zona Compostaje.", "8.3"),
        ("Movimientos", "Modal rediseñado en 3 secciones con filtros, restricciones por stock y unidades visibles.", "8.4"),
        ("Productos", "El stock se muestra siempre con su unidad correspondiente.", "8.5"),
    ]
    t = doc.add_table(rows=1 + len(rows), cols=3)
    # (sin style: aplicamos bordes y sombreado manualmente)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        set_cell_shading(c, COLOR_VERDE_BG)
        add_borders(c, color=COLOR_VERDE_FG, sz="6")
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(COLOR_VERDE_FG)
        run.font.size = Pt(11)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            if i % 2 == 1:
                set_cell_shading(c, "F8FAFC")
            add_borders(c, color="CBD5E1", sz="4")
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(10.5)
    created.append(t._element)

    # Separador y línea final
    p3 = doc.add_paragraph()
    created.append(p3._p)

    return created


banner_elements = build_banner_paragraphs(doc)

# Mover los elementos del banner al inicio (justo antes del primer Heading 1)
if first_h1 is not None:
    for el in banner_elements:
        # Sacarlos de su posición actual
        el.getparent().remove(el)
        # Insertarlos antes del primer H1
        first_h1.addprevious(el)


# =============================================================================
# 2) SECCIÓN 8: NOVEDADES DETALLADAS AL FINAL
# =============================================================================
print("Añadiendo sección 8 con detalle de novedades…")

doc.add_page_break()
doc.add_heading("8. Novedades de esta actualización", level=1)
styled_paragraph(
    doc,
    "Este capítulo recoge en detalle todas las mejoras incorporadas a "
    "ViverApp en la última iteración. Cada apartado describe qué ha cambiado, "
    "cómo afecta al uso diario y qué deberías revisar la primera vez que "
    "abras esa pantalla tras la actualización.",
    size=11, after=8,
)

# ---------------------------------------------------------------------------
# 8.1 Modal de bienvenida
# ---------------------------------------------------------------------------
doc.add_heading("8.1 Modal de bienvenida", level=2)
styled_paragraph(
    doc,
    "Cuando un usuario accede por primera vez al Panel de control (Dashboard), "
    "ViverApp muestra un cuadro de bienvenida que explica de un vistazo qué "
    "ofrece la aplicación y enlaza dos recursos:",
)
add_bullet(doc, "Guía de uso en PDF (descargable).")
add_bullet(doc, "Vídeo explicativo paso a paso.")

styled_paragraph(
    doc,
    "El modal se muestra una sola vez por dispositivo. Si quieres volver a "
    "verlo, pulsa el botón “?” verde situado en la cabecera, junto al botón "
    "Salir. También puedes marcar la casilla “Mostrar al iniciar” dentro "
    "del modal para que vuelva a aparecer en cada entrada al Dashboard.",
)

callout_box(
    doc,
    "💡 Recomendación",
    "Pide a cada nuevo usuario que vea el vídeo explicativo y descargue la guía "
    "PDF la primera vez que entra. Así reduces la curva de aprendizaje y las "
    "consultas básicas al equipo técnico.",
    bg_hex=COLOR_VERDE_BG, fg_hex=COLOR_VERDE_FG,
)

# ---------------------------------------------------------------------------
# 8.2 Categorías con formatos y unidades propias
# ---------------------------------------------------------------------------
doc.add_heading("8.2 Categorías con formatos y unidades propias", level=2)
styled_paragraph(
    doc,
    "Hasta esta versión, todos los productos compartían un único concepto de "
    "“tamaño” y se contaban en unidades. Ahora cada categoría usa su propio "
    "control de formato y la unidad apropiada:",
)

add_kv_table(
    doc,
    header=["Categoría", "Campo de formato", "Unidad cantidad", "Decimales"],
    rows=[
        ("Plantas", "Tamaño: Semillero / M12 / M20 / M35", "ud", "No"),
        ("Fitosanitarios", "Formato: Polvo Seco / Polvo Dispersable / Polvo Soluble / Líquido / Pasta / Granulado", "kg o lt (dinámico)", "Sí (hasta 3 decimales)"),
        ("Fertilizantes", "Formato: igual que fitosanitarios", "kg o lt (dinámico)", "Sí (hasta 3 decimales)"),
        ("Áridos / Material Vegetal", "Sin formato (fijo)", "m³", "No"),
        ("Ferretería (alambre, malla, cinturones)", "Sin formato (fijo)", "m", "No"),
        ("Ferretería (resto)", "Sin formato (fijo)", "ud", "No"),
    ],
    header_bg=COLOR_AMBAR_BG, header_fg=COLOR_AMBAR_FG,
)

callout_box(
    doc,
    "⚠️ Importante",
    "Para Fitosanitarios y Fertilizantes la etiqueta de la cantidad cambia "
    "automáticamente: si seleccionas el formato “Líquido” verás “Litros (lt)”; "
    "para cualquier otro formato verás “Kilogramos (kg)”.",
    bg_hex=COLOR_AMBAR_BG, fg_hex=COLOR_AMBAR_FG,
)

# ---------------------------------------------------------------------------
# 8.3 Nuevas zonas en el mapa del vivero
# ---------------------------------------------------------------------------
doc.add_heading("8.3 Nuevas zonas en el mapa del vivero", level=2)
styled_paragraph(
    doc,
    "El mapa del vivero incorpora cuatro zonas no numéricas, dedicadas a "
    "categorías concretas. Cada zona está pintada en su propio color en "
    "el mapa y aparece tanto en los filtros de la tabla de movimientos "
    "como en los desplegables del modal de Nuevo movimiento.",
)

add_kv_table(
    doc,
    header=["Zona", "Categorías que admite", "Uso típico"],
    rows=[
        ("Almacén Fitosanitarios", "Fitosanitarios", "Almacenamiento exclusivo de productos fitosanitarios (pesticidas, fungicidas, etc.)."),
        ("Almacén General", "Ferretería", "Stock de alambres, mallas, cinturones y resto de material de ferretería."),
        ("Almacén Fertilizantes", "Fertilizantes", "Almacenamiento exclusivo de fertilizantes (sólidos y líquidos)."),
        ("Zona Compostaje", "Áridos / Material Vegetal", "Recepción y acopio de áridos y material vegetal."),
    ],
    header_bg=COLOR_VERDE_BG, header_fg=COLOR_VERDE_FG,
)

styled_paragraph(
    doc,
    "Regla clave: cuando seleccionas un producto en el modal de movimiento, "
    "el desplegable de zona origen y zona destino solo te ofrece la zona "
    "que corresponde a su categoría. Esto evita errores y mezclas de stock.",
    bold=False, size=10.5, after=8,
)

callout_box(
    doc,
    "✅ Comprobaciones tras la actualización",
    "1) Comprueba que la imagen del mapa muestra los tres almacenes y la Zona "
    "Compostaje. 2) Pulsa sobre cada nueva zona y verifica que el modal listará "
    "los productos correspondientes. 3) Si la imagen sigue siendo la antigua, "
    "haz un Ctrl+Shift+R en el navegador para vaciar caché.",
    bg_hex=COLOR_VERDE_BG, fg_hex=COLOR_VERDE_FG,
)

# ---------------------------------------------------------------------------
# 8.4 Modal de movimientos rediseñado
# ---------------------------------------------------------------------------
doc.add_heading("8.4 Modal “Nuevo movimiento” rediseñado", level=2)
styled_paragraph(
    doc,
    "El modal de registro de movimientos se ha reorganizado en tres secciones "
    "visualmente distintas, cada una con su propio tono de color:",
)

add_kv_table(
    doc,
    header=["Sección", "Color", "Qué contiene"],
    rows=[
        ("1. Origen y destino", "Azul", "Origen, destino, zona origen/destino, distrito/barrio/dirección, fecha M35 y préstamo."),
        ("2. Producto", "Verde", "Filtros de categoría, subcategoría y buscador, más el selector del producto."),
        ("3. Detalles del producto", "Ámbar", "Formato/Tamaño, cantidad (con unidad) y observaciones."),
    ],
    header_bg=COLOR_AZUL_BG, header_fg=COLOR_AZUL_FG,
)

# 8.4.1
doc.add_heading("8.4.1 Filtros de categoría y subcategoría", level=3)
styled_paragraph(
    doc,
    "Junto al buscador de producto aparecen dos nuevos desplegables: "
    "“Categoría” y “Subcategoría”. La subcategoría se carga dinámicamente "
    "según la categoría seleccionada. Si eliges directamente un producto "
    "desde el desplegable, los filtros se rellenan automáticamente con la "
    "categoría y subcategoría del producto.",
)
styled_paragraph(
    doc,
    "Si después cambias la categoría o subcategoría y el producto seleccionado "
    "ya no encaja con esos filtros, el producto se resetea y debes volver a "
    "elegir uno compatible.",
)

# 8.4.2
doc.add_heading("8.4.2 Solo productos con stock", level=3)
styled_paragraph(
    doc,
    "Cuando registras una salida o un traslado interno (es decir, origen = "
    "Vivero), el desplegable de producto solo lista los productos que tienen "
    "stock real disponible. Si además ya has elegido zona o tamaño origen, "
    "la lista se acota aún más para mostrar solo lo que hay en esa "
    "combinación.",
)
styled_paragraph(
    doc,
    "El selector de producto queda deshabilitado hasta que rellenas origen y "
    "destino, con un mensaje recordatorio (“Elige primero el origen y el "
    "destino del movimiento”).",
)

# 8.4.3
doc.add_heading("8.4.3 Reset al cambiar origen, destino o producto", level=3)
styled_paragraph(
    doc,
    "Cuando se cambia el origen o el destino del movimiento, todos los campos "
    "específicos del producto (categoría, subcategoría, búsqueda, producto, "
    "cantidad, formato, zonas y distribución multi-zona) se limpian. Esto "
    "evita que un producto del flujo anterior quede asociado al nuevo flujo "
    "por error.",
)
styled_paragraph(
    doc,
    "De forma análoga, cambiar el producto limpia formato, zonas, cantidad y "
    "fecha de disponibilidad — todo lo que dependía del producto anterior.",
)

# 8.4.4
doc.add_heading("8.4.4 Layout compacto Origen / Destino", level=3)
styled_paragraph(
    doc,
    "Origen, Destino, Zona origen (si aplica) y Zona destino (si aplica) "
    "ahora caben en una sola línea con dos, tres o cuatro columnas según el "
    "flujo:",
)
add_bullet(doc, "Salida desde Vivero → 3 columnas: Origen | Zona origen | Destino.")
add_bullet(doc, "Entrada al Vivero → 3 columnas: Origen | Destino | Zona destino.")
add_bullet(doc, "Traslado interno Vivero → Vivero → 4 columnas con ambas zonas.")
add_bullet(doc, "Entrada/salida totalmente externa → 2 columnas: Origen | Destino.")

# 8.4.5
doc.add_heading("8.4.5 Campo Formato unificado", level=3)
styled_paragraph(
    doc,
    "Donde antes existían “Formato origen” y “Formato destino” independientes, "
    "ahora hay un único campo Formato dentro de la sección de Detalles. Si un "
    "producto entra al vivero con cierto formato, sale con el mismo, así que "
    "no tiene sentido permitir cambiarlo en el camino.",
)
styled_paragraph(
    doc,
    "Para categorías con formato fijo (Áridos, Material Vegetal y Ferretería) "
    "el campo Formato no aparece en pantalla: solo se ve la cantidad con su "
    "unidad correspondiente entre paréntesis (“m³”, “m” o “unidades”).",
)

# 8.4.6
doc.add_heading("8.4.6 Cantidades con unidad y decimales", level=3)
styled_paragraph(
    doc,
    "Tanto en la tabla principal de Movimientos como en el modal de detalle "
    '“Ver”, la columna “Cant.” muestra ahora la cantidad acompañada de la '
    "unidad: 5 ud, 2.5 lt, 25 kg, 3 m³, etc.",
)
styled_paragraph(
    doc,
    "Internamente las cantidades admiten hasta 3 decimales (precisión NUMERIC "
    "(12, 3) en base de datos), pero se muestran con un máximo de 2 decimales "
    "y sin ceros sobrantes para no saturar la lectura.",
)

callout_box(
    doc,
    "🧪 Caso típico fitosanitarios",
    "Para registrar “2 botellas de 5 L” de un fitosanitario líquido, selecciona "
    "el formato “Líquido” en la sección de detalles, escribe 10 en el campo "
    "“Litros (lt)” y opcionalmente añade en Observaciones la nota "
    '"2 botellas de 5L". Así el stock cuadra y la trazabilidad queda intacta.',
    bg_hex=COLOR_AMBAR_BG, fg_hex=COLOR_AMBAR_FG,
)

# ---------------------------------------------------------------------------
# 8.5 Tabla de Productos con unidad
# ---------------------------------------------------------------------------
doc.add_heading("8.5 Tabla de Productos con unidad en el stock", level=2)
styled_paragraph(
    doc,
    "La tabla de Productos muestra ahora la unidad junto al valor numérico "
    "de Stock y Stock mínimo. La unidad se infiere de la categoría del "
    "producto:",
)

add_kv_table(
    doc,
    header=["Categoría", "Cómo se ve el stock"],
    rows=[
        ("Plantas", "“5 ud”, “120 ud”, …"),
        ("Fitosanitario / Fertilizante", "“12 kg/lt” (mixto, porque convive kg y lt)"),
        ("Áridos / Material Vegetal", "“3 m³”"),
        ("Ferretería en metros", "“45 m”"),
        ("Ferretería en unidades", "“60 ud”"),
    ],
    header_bg=COLOR_GRIS_BG, header_fg=COLOR_NEGRO,
)

# ---------------------------------------------------------------------------
# 8.6 Otros ajustes y mejoras menores
# ---------------------------------------------------------------------------
doc.add_heading("8.6 Otros ajustes y mejoras", level=2)
styled_paragraph(doc, "Mejoras adicionales que pueden pasar inadvertidas pero que cambian la experiencia diaria:")
add_bullet(
    doc,
    "Los nombres de las zonas especiales se muestran siempre legibles "
    "(“Almacén Fitosanitarios”, “Zona Compostaje”) "
    "aunque internamente se guarden como identificadores cortos "
    "(almacen-fito, compostaje, etc.).",
)
add_bullet(
    doc,
    "Al pinchar una zona del mapa, los items se filtran tolerando tildes, "
    "guiones y mayúsculas en el identificador.",
)
add_bullet(
    doc,
    "Cuando un endpoint del servidor falla por un error interno, ahora "
    "devuelve un mensaje útil en la pantalla en vez de un genérico "
    "“CORS error”. Esto agiliza la resolución de incidencias.",
)
add_bullet(
    doc,
    "Los flujos de pedidos siguen exactamente igual: no se ha cambiado el "
    "ciclo de vida ni los estados (PENDIENTE → APROBADO → SERVIDO).",
)

# Cierre
doc.add_paragraph()
add_section_divider(doc)
styled_paragraph(
    doc,
    "Si detectas comportamientos que no encajan con lo descrito en este "
    "capítulo, recuerda hacer un Ctrl+Shift+R en el navegador para vaciar "
    "caché. Si persiste, contacta con el equipo técnico siguiendo los "
    "canales del capítulo 7.",
    italic := False, size=10.5,
)

# =============================================================================
# GUARDADO
# =============================================================================
DST.parent.mkdir(parents=True, exist_ok=True)
print(f"Guardando {DST}…")
doc.save(str(DST))
print("OK")
